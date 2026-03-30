import os
from docx import Document

def generate_docx(llm_output: str, output_path: str, template_path: str):
    """
    Reads the docx template, and appends the LLM output.
    """
    # Open template if it exists, otherwise create new
    if os.path.exists(template_path):
        try:
            doc = Document(template_path)
            # Strip all dummy content from template but preserve the styles/metadata
            for p in doc.paragraphs:
                p._element.getparent().remove(p._element)
            for t in doc.tables:
                t._element.getparent().remove(t._element)
        except Exception:
            doc = Document()
    else:
        doc = Document()
        
    doc.add_heading("AI Generated Test Scenarios", level=1)
    
    # Very basic inclusion of text. 
    # Proper Markdown to DOCX translation requires pandoc, but we'll use raw insertion for simplicity and reliability.
    for line in llm_output.split('\n'):
        if line.startswith('#'):
            # Count hashes for heading level
            level = min(line.count('#'), 9)
            text = line.replace('#', '').strip()
            if text:
                doc.add_heading(text, level=level)
        else:
            if line.strip():
                doc.add_paragraph(line.strip())
                
    doc.save(output_path)
    return True
